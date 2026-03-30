"""
humantrace_document_extractor.py
Format-aware text extraction for HumanTrace Institutional.

Accepts .docx, .pdf, or plain text and returns clean extracted text
ready to pass into DocumentInput for consistency scoring and BSE matching.

Supported formats:
  .docx  — python-docx (preferred, lossless)
  .pdf   — pdfminer.six (text-layer PDFs) with pytesseract OCR fallback
  .txt   — direct read
  bytes  — raw text bytes (from API upload)

Install dependencies:
  pip install python-docx pdfminer.six
  pip install pytesseract pillow pdf2image   # optional OCR fallback
"""

from __future__ import annotations

import io
import os
import tempfile
from pathlib import Path
from typing import Union


# ---------------------------------------------------------------------------
# Extraction result
# ---------------------------------------------------------------------------

class ExtractionResult:
    def __init__(self, text: str, method: str, warning: str = ""):
        self.text    = text.strip()
        self.method  = method    # "docx" | "pdf_text" | "pdf_ocr" | "plain"
        self.warning = warning   # non-empty if fallback was used or quality is low
        self.ok      = bool(self.text)

    def __repr__(self):
        preview = self.text[:80].replace("\n", " ")
        return f"ExtractionResult(method={self.method!r}, chars={len(self.text)}, preview={preview!r})"


# ---------------------------------------------------------------------------
# .docx extraction
# ---------------------------------------------------------------------------

def _extract_docx(data: bytes) -> ExtractionResult:
    try:
        from docx import Document
    except ImportError:
        return ExtractionResult(
            "", "docx",
            "python-docx not installed. Run: pip install python-docx"
        )
    try:
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs)
        if not text:
            return ExtractionResult("", "docx", "Document appears to contain no text paragraphs.")
        return ExtractionResult(text, "docx")
    except Exception as e:
        return ExtractionResult("", "docx", f"docx extraction failed: {e}")


# ---------------------------------------------------------------------------
# .pdf extraction — text layer first, OCR fallback
# ---------------------------------------------------------------------------

def _extract_pdf_text_layer(data: bytes) -> str:
    """Extract text from PDF text layer using pdfminer.six."""
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(io.BytesIO(data))
        return text or ""
    except ImportError:
        return ""
    except Exception:
        return ""


def _extract_pdf_ocr(data: bytes) -> str:
    """OCR fallback for scanned PDFs using pdf2image + pytesseract."""
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
        from PIL import Image

        # Use Tesseract path if configured (Windows)
        tesseract_path = r"C:\Users\Brian-Owlume\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path

        images = convert_from_bytes(data, dpi=200)
        pages = [pytesseract.image_to_string(img) for img in images]
        return "\n".join(pages)
    except ImportError:
        return ""
    except Exception:
        return ""


def _extract_pdf(data: bytes) -> ExtractionResult:
    # Try text layer first
    text = _extract_pdf_text_layer(data)
    if text and len(text.strip()) > 100:
        return ExtractionResult(text, "pdf_text")

    # Fallback to OCR
    text = _extract_pdf_ocr(data)
    if text and len(text.strip()) > 50:
        return ExtractionResult(
            text, "pdf_ocr",
            "PDF had no text layer — extracted via OCR. Quality may vary."
        )

    return ExtractionResult(
        "", "pdf_text",
        "Could not extract text from PDF. "
        "Ensure pdfminer.six is installed (pip install pdfminer.six) "
        "or that the PDF contains a text layer."
    )


# ---------------------------------------------------------------------------
# Image / screenshot extraction — OCR via Tesseract
# ---------------------------------------------------------------------------

def _extract_image(data: bytes) -> ExtractionResult:
    """OCR an image file (PNG, JPG, WEBP, TIFF) using pytesseract."""
    try:
        import pytesseract
        from PIL import Image

        tesseract_path = r"C:\Users\Brian-Owlume\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path

        img  = Image.open(io.BytesIO(data))
        text = pytesseract.image_to_string(img)
        if text and len(text.strip()) > 20:
            return ExtractionResult(text, "image_ocr")
        return ExtractionResult("", "image_ocr", "OCR returned no readable text from image.")
    except ImportError:
        return ExtractionResult(
            "", "image_ocr",
            "pytesseract or Pillow not installed. Run: pip install pytesseract pillow"
        )
    except Exception as e:
        return ExtractionResult("", "image_ocr", f"Image OCR failed: {e}")


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------

def _extract_plain(data: Union[bytes, str]) -> ExtractionResult:
    if isinstance(data, bytes):
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1", errors="replace")
    else:
        text = data
    return ExtractionResult(text, "plain")


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def extract_text(
    source: Union[bytes, str, Path],
    filename: str = "",
) -> ExtractionResult:
    """
    Extract text from a document source.

    Args:
        source:    bytes (file content), str (plain text), or Path (file path)
        filename:  original filename — used to detect format from extension.
                   If source is a Path, filename is inferred automatically.

    Returns:
        ExtractionResult with .text, .method, .warning, .ok

    Usage:
        # From uploaded bytes
        result = extract_text(file_bytes, filename="employment_letter.docx")

        # From a file path
        result = extract_text(Path("/uploads/business_plan.pdf"))

        # Plain text string
        result = extract_text("I have worked at Meridian...")

        if result.ok:
            doc = DocumentInput("el_001", "employer_authored", result.text, "Employment letter")
    """
    # Handle Path input
    if isinstance(source, Path):
        filename = filename or source.name
        with open(source, "rb") as f:
            source = f.read()

    # Handle plain string input
    if isinstance(source, str):
        return _extract_plain(source)

    # Detect format from filename extension
    ext = Path(filename).suffix.lower() if filename else ""

    if ext == ".docx":
        return _extract_docx(source)
    elif ext == ".pdf":
        return _extract_pdf(source)
    elif ext in (".png", ".jpg", ".jpeg", ".webp", ".tiff", ".tif", ".bmp"):
        return _extract_image(source)
    elif ext in (".txt", ".text", ".md"):
        return _extract_plain(source)
    else:
        # No extension or unknown — try docx first, then pdf, then plain text
        result = _extract_docx(source)
        if result.ok:
            return result
        result = _extract_pdf(source)
        if result.ok:
            return result
        return _extract_plain(source)


# ---------------------------------------------------------------------------
# Dependency check — call at startup to warn about missing packages
# ---------------------------------------------------------------------------

def check_dependencies() -> dict[str, bool]:
    """
    Returns a dict of optional dependency availability.
    Call at API startup and log the result.
    """
    deps = {}
    try:
        import docx
        deps["python-docx"] = True
    except ImportError:
        deps["python-docx"] = False

    try:
        import pdfminer
        deps["pdfminer.six"] = True
    except ImportError:
        deps["pdfminer.six"] = False

    try:
        import pytesseract
        deps["pytesseract"] = True
    except ImportError:
        deps["pytesseract"] = False

    try:
        import pdf2image
        deps["pdf2image"] = True
    except ImportError:
        deps["pdf2image"] = False

    return deps


# ---------------------------------------------------------------------------
# Smoke test
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Dependency check:")
    deps = check_dependencies()
    for name, available in deps.items():
        status = "OK" if available else "MISSING"
        print(f"  {name:<20} {status}")

    print("\nPlain text extraction:")
    result = extract_text("I have worked at Meridian Financial Services for eight years.")
    print(f"  {result}")
    print(f"  Text: {result.text[:60]}")

    print("\nInstall missing dependencies with:")
    missing = [name for name, ok in deps.items() if not ok]
    if missing:
        print(f"  pip install {' '.join(missing)}")
    else:
        print("  All dependencies present.")